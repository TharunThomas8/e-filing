from flask import Flask, render_template, request, abort, jsonify
from docx import Document
from datetime import datetime
import os
import tempfile
import io
import boto3
from botocore.exceptions import NoCredentialsError
from dotenv import load_dotenv
import logging
from typing import Dict, Any, Optional
import zipfile
from dataclasses import dataclass, field

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Import your modules with error handling
try:
    from form_fields import FIELDS
    from s3_bucketHandler import serve_s3_file_as_attachment
except ImportError as e:
    logger.error(f"Failed to import modules: {e}")
    FIELDS = []
    # Fallback function if s3_bucketHandler is not available
    def serve_s3_file_as_attachment(s3_client, bucket, path, filename):
        return jsonify({"error": "S3 handler not available"}), 500

@dataclass
class AppConfig:
    """Application configuration class"""
    bucket_name: str = os.getenv("S3_BUCKET_NAME", "efiling-store")
    template_file: str = "template.docx"
    templates: Dict[str, str] = field(default_factory=lambda: {
        "base-template": "template.docx",
        "docket-template": "docket_template.docx",
        "e-stamping-template": "e_stamping_template.docx",
        "Index-template": "Index_template.docx",
        "notice-to-all-respondants-template": "notice_to_all_respondants_template.docx",
        "process-memo-template": "process_memo_template.docx",
        "vakkalath-template": "vakkalath_template.docx"
    })
    output_prefix: str = "/output/"
    host: str = "0.0.0.0"
    port: int = 5000
    debug: bool = False

class DocumentProcessor:
    """Handles document processing operations"""
    
    def __init__(self, s3_client, bucket_name: str):
        self.s3_client = s3_client
        self.bucket_name = bucket_name
    
    @staticmethod
    def get_custom_datetime_format() -> str:
        """Generate custom datetime format for filenames"""
        return datetime.now().strftime("%d_%m_%YT%H_%M_%S")
    
    @staticmethod
    def get_formatted_current_date() -> str:
        """Get formatted current date with ordinal suffix"""
        def get_ordinal_suffix(day: int) -> str:
            if 11 <= day <= 13:
                return 'th'
            return {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
        
        now = datetime.now()
        day = now.day
        suffix = get_ordinal_suffix(day) + ' day of'
        return f"{day}{suffix} {now.strftime('%B, %Y')}"
    
    @staticmethod
    def convert_date_format(date_str: str) -> str:
        """Convert date from YYYY-MM-DD to DD/MM/YYYY format"""
        try:
            return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d/%m/%Y")
        except (ValueError, TypeError) as e:
            logger.warning(f"Date conversion failed for {date_str}: {e}")
            return date_str or ""

    @staticmethod
    def format_number_indian(number: str) -> str:
        """Convert an integer to a string with commas in Indian number format"""
        if len(number) <= 3:
            return number

        last_three = number[-3:]
        rest = number[:-3]

        # Reverse, group by 2s, reverse again
        rest = rest[::-1]
        grouped = [rest[i:i+2] for i in range(0, len(rest), 2)]
        formatted_rest = ','.join(grouped)[::-1]
        return formatted_rest + ',' + last_three

    def _replace_text_in_docx(self, doc: Document, replacements: Dict[str, str]) -> None:
        """Replace text in document paragraphs and tables while preserving formatting"""
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, replacements)
    
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, replacements)

    def _replace_text_in_paragraph(self, paragraph, replacements: Dict[str, str]) -> None:
        """Replace text in a paragraph while preserving formatting"""
        # Get the full text of the paragraph
        full_text = paragraph.text
        
        # Check if any replacements are needed
        needs_replacement = False
        for old_text in replacements.keys():
            if old_text in full_text:
                needs_replacement = True
                break
        
        if not needs_replacement:
            return
        
        # Perform all replacements on the full text
        modified_text = full_text
        for old_text, new_text in replacements.items():
            modified_text = modified_text.replace(old_text, str(new_text))
        
        # If text hasn't changed, no need to update
        if modified_text == full_text:
            return
        
        # Clear existing runs and add new text with base formatting
        if paragraph.runs:
            # Preserve the formatting of the first run
            first_run = paragraph.runs[0]
            font = first_run.font
            
            # Store formatting properties
            font_name = font.name
            font_size = font.size
            bold = font.bold
            italic = font.italic
            underline = font.underline
            color = font.color.rgb if font.color.rgb else None
            
            # Clear all runs
            paragraph.clear()
            
            # Add new run with preserved formatting
            new_run = paragraph.add_run(modified_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.font.bold = bold
            new_run.font.italic = italic
            new_run.font.underline = underline
            if color:
                new_run.font.color.rgb = color
        else:
            # If no runs exist, just add the text
            paragraph.add_run(modified_text)

    def process_docx(self, template_path: str, replacements: Dict[str, str], output_path: str) -> bool:
        """Process DOCX template with replacements and upload to S3"""
        try:
            if not os.path.exists(template_path):
                logger.error(f"Template file not found: {template_path}")
                return False
        
            doc = Document(template_path)
            self._replace_text_in_docx(doc, replacements)
        
            # Save to in-memory buffer
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
        
            # Upload to S3
            self.s3_client.upload_fileobj(output_stream, self.bucket_name, output_path)
            logger.info(f"Document successfully processed and uploaded: {output_path}")
            return True
        
        except FileNotFoundError as e:
            logger.error(f"Template file not found: {e}")
            return False
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return False

class FormDataProcessor:
    """Handles form data processing and validation"""
    
    @staticmethod
    def validate_form_data(data: Dict[str, Any]) -> bool:
        """Validate required form fields"""
        if not FIELDS:
            return True
            
        required_fields = [field["name"] for field in FIELDS if field.get("required", False)]
        missing_fields = [field for field in required_fields if not data.get(field)]
        
        if missing_fields:
            logger.warning(f"Missing required fields: {missing_fields}")
            return False
        return True
    
    @staticmethod
    def build_replacements(data: Dict[str, Any]) -> Dict[str, str]:
        """Build replacement dictionary from form data"""
        replacements = {}
        
        # Process basic field replacements
        for field in FIELDS:
            field_name = field["name"]
            placeholder = field["placeholder"]
            field_value = data.get(field_name, "")
            field_type = field.get("datatype")
            
            if field_type.strip() == "date":
                field_value = DocumentProcessor.convert_date_format(field_value)
            if field_type.strip() == "number":
                field_value = DocumentProcessor.format_number_indian(field_value)
            replacements[placeholder] = str(field_value)
        
        # Add computed fields
        replacements.update(FormDataProcessor._build_computed_replacements(replacements))
        
        return replacements
    
    @staticmethod
    def _build_computed_replacements(replacements: Dict[str, str]) -> Dict[str, str]:
        """Build computed replacement fields"""
        computed = {}
        
        # Build cause of action text
        cause_of_action = f"The cause of action for this claim petition arose on {replacements.get('(DATE1)', '')}, when the respondent initiated proceedings to draw electrical lines through the property of the petitioner, and thereafter on {replacements.get('(COA_DATE1)', '')}, when the respondent paid an amount of ₹{replacements.get('(COA_AMNT1)', '')} as compensation to the petitioner."
        
        # Add jurisdiction clause
        village = replacements.get("(VILLAGE)", "")
        taluk = replacements.get("(TALUK)", "")
        cause_of_action += f"Thereafter, continuously at {village} Village in {taluk} Taluk which is within the jurisdiction of this Honourable Court."
        
        computed["(CAUSE_OF_ACTION)"] = cause_of_action
        
        # Handle ARES2 field
        ares2 = replacements.get("(ARES2)", "0")
        syno2 = replacements.get("(SYNO2)", "")
        if ares2 and ares2 != "0":
            computed["(ARES2)"] = f"and {ares2} of property comprised in Sy.No. {syno2}"
        else:
            computed["(ARES2)"] = ""
        
        # Add other computed fields
        computed["(DISTRICT)"] = replacements.get("(DISTRICT)", "").upper()
        computed["(CURRENT_DATE)"] = DocumentProcessor.get_formatted_current_date()
        
        # Calculate total amount
        try:
            amount1 = int((replacements.get("(AMNT1)", "0") or "0").replace(",", ""))
            amount2 = int((replacements.get("(AMNT2)", "0") or "0").replace(",", ""))
            amount3 = int((replacements.get("(AMNT3)", "0") or "0").replace(",", ""))
            computed["(TOTAL_AMOUNT)"] = DocumentProcessor.format_number_indian(str(amount1 + amount2 + amount3))
        except (ValueError, TypeError):
            computed["(TOTAL_AMOUNT)"] = "0"
            logger.warning("Failed to calculate total amount")
        
        computed["(ESTABLISHMENT)"] = replacements.get("(ESTABLISHMENT)").upper()
        
        return computed

# Initialize Flask app at module level (CRITICAL for Vercel)
app = Flask(__name__)

# Initialize configuration
config = AppConfig()

# Initialize S3 client with error handling
s3_client = None
try:
    s3_client = boto3.client(
        's3',
        aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
        aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
        region_name=os.getenv('AWS_REGION', 'us-east-1')
    )
    logger.info("S3 client initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize S3 client: {e}")

# Initialize processors
doc_processor = None
form_processor = FormDataProcessor()

if s3_client:
    doc_processor = DocumentProcessor(s3_client, config.bucket_name)

@app.route("/", methods=["GET"])
def form():
    """Main page render"""
    try:
        if request.method == "GET":
            if not FIELDS:
                return jsonify({"error": "Form fields not configured"}), 500
            return render_template("form.html", fields=FIELDS)

    except Exception as e:
        logger.error(f"Form processing error: {e}")
        return jsonify({"error": "Internal server error", "details": str(e)}), 500

@app.route("/download-document/<doc_type>", methods=["POST"])
def download_document(doc_type):
    """Handle document download"""
    try:
        # Check if services are available
        if not s3_client or not doc_processor:
            return jsonify({"error": "Service temporarily unavailable"}), 503
            
        # Check if document type exists
        if doc_type not in config.templates:
            return jsonify({"error": f"Document type '{doc_type}' not found"}), 404
            
        form_data = request.form.to_dict()
        logger.info(f"Processing document download for type: {doc_type}")
        
        # Validate form data
        if not form_processor.validate_form_data(form_data):
            return jsonify({"error": "Missing required form fields"}), 400
        
        
        replacements = form_processor.build_replacements(form_data)
        
        # Handle petitioner address
        if request.form.get('petitioner_address_checker') == 'on':
            village = replacements.get("(VILLAGE)", "")
            taluk = replacements.get("(TALUK)", "")
            district = replacements.get("(DISTRICT)", "")
            pincode = replacements.get("(PINCODE)", "")
            replacements["(PETITIONER_ADDRESS)"] = f"{village} Village, {taluk} Taluk, {district} District. PIN -{pincode}"
        else:
            replacements["(PETITIONER_ADDRESS)"] = ""
        
        # Generate output filename and path for document
        output_filename = f"{doc_processor.get_custom_datetime_format()}_{doc_type}_output.docx"
        output_path = f"{config.output_prefix}{output_filename}"
        
        # Process document
        template_path = config.templates[doc_type]
        if not doc_processor.process_docx(template_path, replacements, output_path):
            return jsonify({"error": "Document processing failed"}), 500
        
        # Serve file
        return serve_s3_file_as_attachment(s3_client, config.bucket_name, output_path, output_filename)
        
    except Exception as e:
        logger.error(f"Document processing error: {e}")
        return jsonify({"error": "Internal server error", "details": str(e)}), 500

# Download all from S3 for serving the zip file
@app.route("/download-all-documents-s3", methods=["POST"])
def download_all_documents_s3():
    """Handle download of all documents as a zip file (served from S3)"""
    try:
        # Check if services are available
        if not s3_client or not doc_processor:
            return jsonify({"error": "Service temporarily unavailable"}), 503
        
        form_data = request.form.to_dict()
        logger.info("Processing download all documents request (S3 version)")
        
        # Validate form data
        if not form_processor.validate_form_data(form_data):
            return jsonify({"error": "Missing required form fields"}), 400
        
        # Build replacements once for all documents
        replacements = form_processor.build_replacements(form_data)
        
        # Handle petitioner address
        if request.form.get('petitioner_address_checker') == 'on':
            village = replacements.get("(VILLAGE)", "")
            taluk = replacements.get("(TALUK)", "")
            district = replacements.get("(DISTRICT)", "")
            pincode = replacements.get("(PINCODE)", "")
            replacements["(PETITIONER_ADDRESS)"] = f"{village} Village, {taluk} Taluk, {district} District. PIN -{pincode}"
        else:
            replacements["(PETITIONER_ADDRESS)"] = ""
        
        # Create in-memory zip file
        zip_buffer = io.BytesIO()
        datetime_stamp = doc_processor.get_custom_datetime_format()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            processed_docs = []
            failed_docs = []
            
            # Process each document type
            for doc_type, template_path in config.templates.items():
                try:
                    # Generate output filename and path for this document
                    output_filename = f"{datetime_stamp}_{doc_type}_output.docx"
                    output_path = f"{config.output_prefix}{output_filename}"
                    
                    # Process document
                    if doc_processor.process_docx(template_path, replacements, output_path):
                        # Download the processed file from S3
                        try:
                            response = s3_client.get_object(Bucket=config.bucket_name, Key=output_path)
                            file_content = response['Body'].read()
                            
                            # Add to zip file
                            zip_file.writestr(output_filename, file_content)
                            processed_docs.append(doc_type)
                            logger.info(f"Added {doc_type} to zip file")
                            
                        except Exception as s3_error:
                            logger.error(f"Failed to retrieve {doc_type} from S3: {s3_error}")
                            failed_docs.append(doc_type)
                    else:
                        logger.error(f"Failed to process document type: {doc_type}")
                        failed_docs.append(doc_type)
                        
                except Exception as doc_error:
                    logger.error(f"Error processing {doc_type}: {doc_error}")
                    failed_docs.append(doc_type)
        
        # Check if any documents were processed
        if not processed_docs:
            return jsonify({"error": "Failed to process any documents"}), 500
        
        # Upload zip file to S3
        zip_filename = f"{datetime_stamp}_all_documents.zip"
        zip_s3_path = f"{config.output_prefix}{zip_filename}"
        
        s3_client.put_object(
            Bucket=config.bucket_name,
            Key=zip_s3_path,
            Body=zip_buffer.getvalue(),
            ContentType='application/zip'
        )
        
        # Serve zip file from S3
        return serve_s3_file_as_attachment(s3_client, config.bucket_name, zip_s3_path, zip_filename)
        
    except Exception as e:
        logger.error(f"Download all documents error: {e}")
        return jsonify({"error": "Internal server error", "details": str(e)}), 500

@app.route("/health")
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy",
        "s3_available": s3_client is not None,
        "fields_loaded": len(FIELDS) > 0,
        "templates_configured": len(config.templates),
        "timestamp": datetime.now().isoformat()
    })

@app.route("/debug")
def debug_info():
    """Debug information endpoint"""
    return jsonify({
        "environment_variables": {
            "AWS_ACCESS_KEY_ID": "***" if os.getenv('AWS_ACCESS_KEY_ID') else "Not set",
            "AWS_SECRET_ACCESS_KEY": "***" if os.getenv('AWS_SECRET_ACCESS_KEY') else "Not set",
            "AWS_REGION": os.getenv('AWS_REGION', 'Not set'),
            "S3_BUCKET_NAME": os.getenv('S3_BUCKET_NAME', 'Not set')
        },
        "file_structure": {
            "files_in_directory": os.listdir('.'),
            "template_file_exists": os.path.exists(config.template_file),
            "template_files": {name: os.path.exists(path) for name, path in config.templates.items()}
        },
        "services": {
            "s3_client": s3_client is not None,
            "doc_processor": doc_processor is not None,
            "fields_count": len(FIELDS)
        }
    })

@app.errorhandler(400)
def bad_request(error):
    return jsonify({"error": "Bad Request", "message": str(error.description)}), 400

@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Not Found", "message": "The requested resource was not found"}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({"error": "Internal Server Error", "message": "Something went wrong"}), 500

# For local development
if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=5000)